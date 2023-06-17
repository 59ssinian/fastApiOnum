import requests
from docx import Document




'''
# 새로운 워드 문서 생성
document = Document()

# 문서에 단락 추가
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)


records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

# 워드 문서 저장
document.save('example.docx')
'''

import requests
from requests.exceptions import JSONDecodeError


# SharePoint 사이트 URL 설정
site_url = "https://omipc.sharepoint.com/:f:/s/omipc2/Erk9wOt1AndEo0AcPbEaCvABriaosJ5wRKlEZjfHeBlRwQ?e=c0V9pG"
file_path = '/tmp/pycharm_project_410/src/letters/example.docx'

# 저장할 파일 내용
file_content = 'This is the content of my file.'

# 요청 디제스트 값을 가져오기 위한 요청 보내기
api_url = site_url + "/_api/contextinfo"
headers = {
    "accept": "application/json;odata=verbose",
    "content-type": "application/json;odata=verbose"
}
response = requests.post(api_url, headers=headers)


try:
    response.raise_for_status()
    data = response.json()
    request_digest = data['d']['GetContextWebInformation']['FormDigestValue']
except JSONDecodeError:
    print("Error: Unable to decode the response content as JSON.")
except requests.exceptions.HTTPError as http_error:
    print(f"HTTP Error: {http_error}")
except KeyError:
    print("Error: Invalid JSON response format. Required data is missing.")
else:
    # 요청 디제스트 값을 사용하여 파일 업로드 요청 보내기
    # (앞서 제공한 파일 업로드 예시 코드와 연결하여 사용)
    api_url = site_url + "/_api/web/getfolderbyserverrelativeurl('" + file_path + "')/files/add(url='your-file.txt', overwrite=true)"
    headers = {
        "accept": "application/json;odata=verbose",
        "content-type": "application/json;odata=verbose",
        "X-RequestDigest": request_digest
    }
    data = {
        'content': file_content
    }
    response = requests.post(api_url, headers=headers, json=data)


# 요청 결과 확인
if response.status_code == 200:
    print("파일이 성공적으로 업로드되었습니다.")
else:
    print("파일 업로드에 실패했습니다. 상태 코드:", response.status_code)




